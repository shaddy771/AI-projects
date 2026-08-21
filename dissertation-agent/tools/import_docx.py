#!/usr/bin/env python3
"""Import a .docx dissertation draft into Markdown for the dissertation agent."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Нужен пакет python-docx. Установите: pip install -r dissertation-agent/requirements.txt"
    ) from exc


HEADING_RE = re.compile(r"^heading\s*(\d+)$", re.I)
CHAPTER_RE = re.compile(
    r"^(введение|заключение|список\s+литературы|приложение|"
    r"глава\s+\d+|chapter\s+\d+)",
    re.I,
)


def style_level(paragraph) -> int | None:
    style = paragraph.style
    if style is None:
        return None
    name = (style.name or "").strip()
    match = HEADING_RE.match(name)
    if match:
        return int(match.group(1))
    # Outline level from Word, if present
    try:
        level = paragraph._element.pPr.outlineLvl.val  # type: ignore[attr-defined]
        return int(level) + 1
    except Exception:
        return None


def is_list_item(paragraph) -> bool:
    try:
        return paragraph._element.pPr.numPr is not None  # type: ignore[union-attr]
    except Exception:
        return False


def paragraph_to_markdown(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    level = style_level(paragraph)
    if level:
        return f"{'#' * min(level, 6)} {text}"

    if is_list_item(paragraph):
        return f"- {text}"

    return text


def extract_tables(document) -> list[str]:
    blocks: list[str] = []
    for idx, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        width = max(len(r) for r in rows)
        normalized = [r + [""] * (width - len(r)) for r in rows]
        header = normalized[0]
        lines = [
            f"### Таблица {idx}",
            "",
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        blocks.append("\n".join(lines))
    return blocks


def convert_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = [
        f"# Импорт из `{path.name}`",
        "",
        "> Черновик сконвертирован из DOCX. Проверьте заголовки и таблицы перед правкой.",
        "",
    ]

    for paragraph in document.paragraphs:
        md = paragraph_to_markdown(paragraph)
        if md:
            parts.append(md)
            parts.append("")

    tables = extract_tables(document)
    if tables:
        parts.append("## Таблицы из документа")
        parts.append("")
        parts.extend(tables)

    return "\n".join(parts).rstrip() + "\n"


def slugify(title: str) -> str:
    cleaned = re.sub(r"[^\w\s\-а-яА-ЯёЁ]+", "", title, flags=re.UNICODE)
    cleaned = re.sub(r"\s+", "-", cleaned.strip().lower())
    return cleaned[:60] or "section"


def split_by_top_headings(markdown: str) -> dict[str, str]:
    lines = markdown.splitlines()
    sections: dict[str, list[str]] = {}
    preamble: list[str] = []
    current: str | None = None

    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            if CHAPTER_RE.search(title):
                current = slugify(title)
                sections.setdefault(current, []).append(line)
                continue
        if current is None:
            preamble.append(line)
        else:
            sections[current].append(line)

    result: dict[str, str] = {}
    notice = "\n".join(preamble).strip()
    for name, body in sections.items():
        chunk = "\n".join(body).strip()
        if notice:
            chunk = f"{notice}\n\n{chunk}"
        if chunk:
            result[name] = chunk + "\n"

    if not result:
        result["00-full-import"] = markdown if markdown.endswith("\n") else markdown + "\n"
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description="Конвертация DOCX диссертации в Markdown")
    parser.add_argument("docx", type=Path, help="Путь к .docx файлу")
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        default=Path("dissertation-agent/workspace/drafts"),
        help="Куда сохранить Markdown",
    )
    parser.add_argument(
        "--split",
        action="store_true",
        help="Пытаться разбить на файлы по главам/введению/заключению",
    )
    args = parser.parse_args()

    if not args.docx.exists():
        print(f"Файл не найден: {args.docx}", file=sys.stderr)
        return 1
    if args.docx.suffix.lower() != ".docx":
        print("Ожидается файл .docx", file=sys.stderr)
        return 1

    args.output_dir.mkdir(parents=True, exist_ok=True)
    markdown = convert_docx(args.docx)

    if args.split:
        sections = split_by_top_headings(markdown)
        written = []
        for name, body in sections.items():
            out = args.output_dir / f"{name}.md"
            out.write_text(body, encoding="utf-8")
            written.append(out)
        print(f"Сохранено файлов: {len(written)}")
        for path in written:
            print(f" - {path}")
    else:
        out = args.output_dir / f"{args.docx.stem}.md"
        out.write_text(markdown, encoding="utf-8")
        print(f"Сохранено: {out}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
